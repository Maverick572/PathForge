# file_parser.py  v4.1
# Extracts plain text from base64-encoded resume files.
# v4.1: debug logging added.

import base64
import io


def _log(level: str, fn: str, msg: str) -> None:
    print(f"[file_parser:{fn}] {level:<5}  {msg}")

def _info(fn, msg):  _log("INFO",  fn, msg)
def _warn(fn, msg):  _log("WARN",  fn, msg)
def _debug(fn, msg): _log("DEBUG", fn, msg)


def extract_text(b64: str, filename: str) -> str:
    fn = "extract_text"
    _info(fn, f"File: '{filename}', b64 length: {len(b64)}")

    if "," in b64:
        b64 = b64.split(",", 1)[1]
        _debug(fn, "Stripped data-URI prefix")

    try:
        raw = base64.b64decode(b64)
        _info(fn, f"Decoded {len(raw)} bytes")
    except Exception as e:
        _warn(fn, f"base64 decode failed: {e}")
        return ""

    fname = filename.lower()

    if fname.endswith(".pdf"):
        _info(fn, "Attempting PDF extraction")
        text = _extract_pdf(raw)
        if text:
            return text
        _warn(fn, "PDF extraction returned empty — trying plain text fallback")

    if fname.endswith(".docx"):
        _info(fn, "Attempting DOCX extraction")
        text = _extract_docx(raw)
        if text:
            return text
        _warn(fn, "DOCX extraction returned empty — trying plain text fallback")

    _info(fn, "Trying plain text fallback")
    return _extract_plaintext(raw)


def _extract_pdf(raw: bytes) -> str:
    fn = "_extract_pdf"
    try:
        import fitz
        doc   = fitz.open(stream=raw, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            t = page.get_text()
            _debug(fn, f"Page {i+1}: {len(t)} chars")
            pages.append(t)
        text = "\n".join(pages)
        if text.strip():
            _info(fn, f"PDF extracted successfully: {len(text)} chars, {len(doc)} pages")
            return text
        _warn(fn, "PDF parsed but all pages returned empty text — likely image-based/scanned")
    except ImportError:
        _warn(fn, "PyMuPDF not installed — pip install pymupdf")
    except Exception as e:
        _warn(fn, f"PDF extraction error: {e}")
    return ""


def _extract_docx(raw: bytes) -> str:
    fn = "_extract_docx"
    try:
        import docx
        doc    = docx.Document(io.BytesIO(raw))
        parts  = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
                        tables += 1
        text = "\n".join(parts)
        _info(fn, f"DOCX extracted: {len(text)} chars, {len(parts)} paragraphs, {tables} table cells")
        return text
    except ImportError:
        _warn(fn, "python-docx not installed — pip install python-docx")
    except Exception as e:
        _warn(fn, f"DOCX extraction error: {e}")
    return ""


def _extract_plaintext(raw: bytes) -> str:
    fn = "_extract_plaintext"
    try:
        text = raw.decode("utf-8", errors="ignore")
        word_count = text.count(" ")
        _debug(fn, f"Plain text decode: {len(text)} chars, ~{word_count} spaces")
        if word_count > 20:
            _info(fn, f"Plain text fallback succeeded: {len(text)} chars")
            return text
        _warn(fn, f"Plain text has too few spaces ({word_count}) — likely binary, not text")
    except Exception as e:
        _warn(fn, f"Plain text fallback failed: {e}")
    return ""